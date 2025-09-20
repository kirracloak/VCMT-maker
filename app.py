# app.py
import io
import re
import datetime
import streamlit as st
from typing import List, Dict, Iterable, Tuple
from docx import Document
from docx.shared import Inches

# ---------- helpers ----------
def normalise_space(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "")).strip()

def all_doc_text_lines(doc: Document) -> Iterable[str]:
    for p in doc.paragraphs:
        if p.text:
            yield p.text
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text:
                        yield p.text

def load_docx(file_bytes: bytes) -> Document:
    return Document(io.BytesIO(file_bytes))

def extract_units_from_doc(doc: Document) -> Dict[str, Dict]:
    paras = [normalise_space(t) for t in all_doc_text_lines(doc) if normalise_space(t)]
    full_text = "\n".join(paras)
    codes = sorted(set(re.findall(r"\b[A-Z]{3,}[A-Z0-9]{2,}\b", full_text)))
    units: Dict[str, Dict] = {}
    for code in codes:
        name = ""
        for i, line in enumerate(paras):
            if code in line:
                m = re.search(rf"{code}\s*[-:]\s*(.+)", line)
                if m:
                    name = normalise_space(m.group(1))
                elif i + 1 < len(paras) and len(paras[i + 1].split()) >= 3:
                    name = paras[i + 1]
                break
        units[code] = {"code": code, "name": name}
    return units

def list_tables_info(doc: Document) -> List[str]:
    out = []
    for i, t in enumerate(doc.tables):
        header = []
        if len(t.rows) > 0:
            header = [normalise_space(c.text) for c in t.rows[0].cells]
        out.append(f"Table {i}: {len(t.rows)} rows x {len(t.columns)} cols | header: {' | '.join(header[:6]) if header else '(none)'}")
    return out

def find_part1_table_index(doc: Document) -> int:
    """Heuristic: return first table with at least 4 columns (assumed Part 1 table)."""
    for i, t in enumerate(doc.tables):
        if len(t.columns) >= 4:
            return i
    return -1

def get_section_text_for_unit(doc: Document, unit_code: str) -> Tuple[str, List[str]]:
    """
    Heuristic extraction of Application Statement and Performance Evidence.
    Returns (application_short_excerpt, performance_evidence_bullets)
    """
    paras = [normalise_space(t) for t in all_doc_text_lines(doc) if normalise_space(t)]
    # find near lines containing unit_code
    idxs = [i for i, p in enumerate(paras) if unit_code in p]
    app_excerpt = ""
    perf_bullets: List[str] = []
    # search around first occurrence
    if idxs:
        start = max(0, idxs[0] - 6)
        window = paras[start:start + 20]
    else:
        window = paras[:40]

    # find Application Statement
    for i, line in enumerate(window):
        if re.search(r"Application Statement|Application of the unit|Application of skill", line, re.IGNORECASE):
            # take following 1-3 lines that look like sentence(s)
            for j in range(i+1, min(i+4, len(window))):
                if len(window[j].split()) >= 4:
                    app_excerpt = window[j]
                    break
            break
    # fallback: attempt to extract a short sentence near code
    if not app_excerpt:
        for line in window:
            if len(line.split()) >= 6 and "performance" not in line.lower():
                app_excerpt = line
                break

    # find Performance Evidence section lines in the doc (global search)
    all_paras = [normalise_space(t) for t in all_doc_text_lines(doc) if normalise_space(t)]
    for i, p in enumerate(all_paras):
        if re.search(r"Performance Evidence", p, re.IGNORECASE):
            # collect subsequent bullets (lines starting with • - or numbers) up to 20 lines
            for j in range(i+1, min(i+20, len(all_paras))):
                line = all_paras[j]
                if re.match(r"^\s*(?:•|-|\d+\.)\s*", line):
                    perf_bullets.append(re.sub(r"^\s*(?:•|-|\d+\.)\s*", "", line))
                elif len(line.split()) > 4:
                    # some templates use plain sentences; collect a few
                    perf_bullets.append(line)
                else:
                    # stop if we hit another heading-like short line
                    if len(line.split()) <= 3:
                        break
            break

    # keep unique, trimmed and short list
    perf_bullets = [normalise_space(x) for x in perf_bullets][:12]
    return (normalise_space(app_excerpt), perf_bullets)

def mask_evidence_id(eid: str) -> str:
    eid = (eid or "").strip()
    if not eid or eid.lower() == "pending":
        return "Pending"
    if len(eid) <= 4:
        return "*" * (len(eid)-1) + eid[-1] if len(eid) > 1 else eid
    return "*" * (len(eid)-4) + eid[-4:]

def validate_year(y: str) -> bool:
    try:
        yr = int(y)
        this_year = datetime.date.today().year
        return 1000 <= yr <= this_year
    except:
        return False

# ---------- Streamlit UI ----------
st.set_page_config(page_title="VCMT Unit Code Reader (Lite) + Filler", page_icon="🗂", layout="wide")
st.title("VCMT Unit Code Reader (Lite) — Fill & Export")

st.header("Step 1 — Upload VCMT Template (.docx)")
uploaded = st.file_uploader("Upload the Autodocs VCMT template", type=["docx"])
if not uploaded:
    st.info("Upload the VCMT .docx template to begin.")
    st.stop()

try:
    doc = load_docx(uploaded.read())
    st.success("Template loaded.")
except Exception as e:
    st.error(f"Could not load .docx: {e}")
    st.stop()

with st.expander("Template tables found", expanded=False):
    for line in list_tables_info(doc):
        st.write(line)

st.header("Step 2 — Detect & select unit(s)")
extracted_units = extract_units_from_doc(doc)
discovered_codes = sorted(extracted_units.keys())
st.caption("Codes detected from both paragraphs and table cells.")

user_unit_codes = st.multiselect(
    "Choose unit code(s) to complete (enter manually if not visible).",
    options=discovered_codes,
    default=[]
)

manual_units_input = st.text_input("Optional: add unit code(s) comma-separated (e.g., BSBWHS211, SITXWHS005)")
if manual_units_input.strip():
    added = [normalise_space(x).upper() for x in manual_units_input.split(",") if normalise_space(x)]
    user_unit_codes = list(dict.fromkeys([u.upper() for u in user_unit_codes] + added))

# validate against ALL text (case-insensitive)
full_text = "\n".join(all_doc_text_lines(doc))
full_text_up = full_text.upper()
validated_codes = [c for c in user_unit_codes if c.strip().upper() in full_text_up]
invalid_codes = [c for c in set(user_unit_codes) - set(validated_codes)]

if invalid_codes:
    st.warning("The following codes were not found in document text and will be skipped: " + ", ".join(invalid_codes))

if not validated_codes:
    st.info("Please select at least one valid unit code.")
    st.stop()

validated_codes = [c.strip().upper() for c in validated_codes]
st.success("Validated unit codes: " + ", ".join(validated_codes))

# session storage for collected data per unit
if "units_data" not in st.session_state:
    st.session_state.units_data = {}

part1_table_index = find_part1_table_index(doc)
if part1_table_index == -1:
    st.warning("No table with >=4 columns found. The app will still collect data but cannot auto-insert rows into the template. You can download the filled file as a new docx with an appended table instead.")

# Loop through each unit and collect data
st.header("Step 3 — For each unit, answer prompts")
for unit_code in validated_codes:
    unit_name = extracted_units.get(unit_code, {}).get("name", "")
    key = f"unit__{unit_code}"
    if key not in st.session_state.units_data:
        st.session_state.units_data[key] = {
            "unit_code": unit_code,
            "unit_name": unit_name,
            "part1": [],   # list of dicts: qual_name, year, evidence_id, generated_statement
            "part2": None, # dict role_title, years_exp, evidence_id, generated_statement
            "part3": None, # dict pd_title, year, evidence_id, generated_statement
            "confirmed": False
        }

    st.subheader(f"Unit: {unit_code} — {unit_name}")
    with st.expander(f"Open prompts for {unit_code}", expanded=not st.session_state.units_data[key]["confirmed"]):
        # retrieve heuristic source info
        app_excerpt, perf_bullets = get_section_text_for_unit(doc, unit_code)

        st.markdown("**Unit context (extracted from template)**")
        if app_excerpt:
            st.write(f"Application excerpt: {app_excerpt}")
        else:
            st.write("_No clear application statement found in template._")
        if perf_bullets:
            st.write("Performance Evidence (sample lines):")
            for b in perf_bullets[:8]:
                st.write(f"- {b}")
        else:
            st.write("_No Performance Evidence lines found._")

        # ---------- Part 1 ----------
        st.markdown("### Part 1 — Qualifications / Units of Competency")
        st.write("Enter qualifications/units that support competency. One per line. After entering, press 'Parse Qualifications' to convert into rows.")
        qual_raw = st.text_area(f"Qualifications for {unit_code} (format: one per line)", key=f"qual_raw_{unit_code}", value="")
        if st.button(f"Parse Qualifications for {unit_code}", key=f"parse_q_{unit_code}"):
            quals = [normalise_space(q) for q in qual_raw.splitlines() if normalise_space(q)]
            parsed = []
            for i, q in enumerate(quals):
                # create placeholders for year and evidence id
                parsed.append({"qual_name": q, "year": "", "evidence_id": ""})
            st.session_state.units_data[key]["part1"] = parsed

        # show current part1 entries and let user fill year & evidence id
        if st.session_state.units_data[key]["part1"]:
            st.write("Now provide Year (YYYY) and People@TAFE Evidence ID (or 'Pending') for each qualification.")
            for idx, entry in enumerate(st.session_state.units_data[key]["part1"]):
                cols = st.columns([4,1,2])
                with cols[0]:
                    st.write(f"**{idx+1}. {entry['qual_name']}**")
                with cols[1]:
                    year_val = st.text_input(f"Year for #{idx+1}", key=f"{unit_code}_p1_year_{idx}", value=entry.get("year",""))
                with cols[2]:
                    eid_val = st.text_input(f"Evidence ID for #{idx+1}", key=f"{unit_code}_p1_eid_{idx}", value=entry.get("evidence_id",""))
                # simple validation and save
                if year_val and not validate_year(year_val):
                    st.error("Enter a valid 4-digit year ≤ current year.")
                st.session_state.units_data[key]["part1"][idx]["year"] = year_val
                st.session_state.units_data[key]["part1"][idx]["evidence_id"] = eid_val

            # auto-generate Column 3 for each qualification by comparing perf bullets heuristically
            for idx, entry in enumerate(st.session_state.units_data[key]["part1"]):
                # very simple matching: common words between perf bullets and qual name
                bullets = perf_bullets or []
                matched = []
                qual_words = set(w.lower() for w in re.findall(r"\w{4,}", entry["qual_name"]))
                for b in bullets:
                    b_words = set(w.lower() for w in re.findall(r"\w{4,}", b))
                    if qual_words & b_words:
                        matched.append(b)
                # fallback: keep first 3 perf bullets
                if not matched and bullets:
                    matched = bullets[:3]
                col3 = f"Within this qualification, I was required to demonstrate competency in the skills and knowledge required to {app_excerpt or unit_name}.\n\nSpecifically relevant were the following course components:\n"
                for m in matched:
                    col3 += f"• {m}\n"
                st.session_state.units_data[key]["part1"][idx]["generated_statement"] = col3

        # ---------- Part 2 ----------
        st.markdown("### Part 2 — Industry / Community Experience")
        p2 = st.session_state.units_data[key].get("part2") or {"role_title":"","years_exp":"","evidence_id":"","generated_statement":""}
        p2_role = st.text_input(f"Role title for {unit_code} (e.g., Plant Operator)", key=f"{unit_code}_p2_role", value=p2.get("role_title",""))
        p2_employer = st.text_input(f"Employer (optional)", key=f"{unit_code}_p2_employer", value=p2.get("employer",""))
        p2_years = st.text_input(f"How many years in this role?", key=f"{unit_code}_p2_years", value=p2.get("years_exp",""))
        p2_eid = st.text_input(f"People@TAFE Evidence ID for role (or Pending)", key=f"{unit_code}_p2_eid", value=p2.get("evidence_id",""))
        # generate responsibilities by matching keywords from perf bullets
        responsibilities = []
        if perf_bullets:
            # sample keywords
            sample = perf_bullets[:6]
            for s in sample:
                responsibilities.append(s)
        else:
            responsibilities = [
                "Conduct pre-start checks and operate equipment under WHS procedures and SOPs.",
                "Identify hazards and apply risk controls using site procedures.",
                "Maintain records to meet compliance and traceability standards."
            ]
        p2_col3 = f"Key responsibilities and tasks relevant to the performance criteria for {unit_code} {unit_name}:\n"
        for r in responsibilities[:7]:
            p2_col3 += f"• {r}\n"
        st.session_state.units_data[key]["part2"] = {"role_title": p2_role + (f" ({p2_employer})" if p2_employer else ""), "years_exp": p2_years, "evidence_id": p2_eid, "generated_statement": p2_col3}

        # ---------- Part 3 ----------
        st.markdown("### Part 3 — Professional Development (PD)")
        p3_title = st.text_input(f"PD title for {unit_code}", key=f"{unit_code}_p3_title", value=(st.session_state.units_data[key].get("part3") or {}).get("pd_title",""))
        p3_year = st.text_input(f"Year completed (YYYY)", key=f"{unit_code}_p3_year", value=(st.session_state.units_data[key].get("part3") or {}).get("year",""))
        p3_eid = st.text_input(f"People@TAFE Evidence ID (or Pending)", key=f"{unit_code}_p3_eid", value=(st.session_state.units_data[key].get("part3") or {}).get("evidence_id",""))
        # generate PD alignment bullets (simple keywords)
        pd_bullets = []
        if perf_bullets:
            pd_bullets = [perf_bullets[i] for i in range(min(4, len(perf_bullets)))]
        else:
            pd_bullets = [
                "Reinforced safe work procedures and hazard identification.",
                "Improved ability to apply risk controls and conduct pre-start checks."
            ]
        p3_col3 = f"This professional development enhanced my ability to meet the performance criteria for {unit_code} {unit_name}. Specifically, it:\n"
        for b in pd_bullets[:4]:
            p3_col3 += f"• {b}\n"
        st.session_state.units_data[key]["part3"] = {"pd_title": p3_title, "year": p3_year, "evidence_id": p3_eid, "generated_statement": p3_col3}

        # confirm block for this unit
        st.markdown("#### Review & confirm entries for this unit")
        def show_unit_summary(u):
            st.write(f"**Unit:** {u['unit_code']} — {u['unit_name']}")
            st.write("**Part 1 entries (Qualifications):**")
            if u["part1"]:
                for i, e in enumerate(u["part1"], start=1):
                    st.write(f"- {i}. {e['qual_name']} | Year: {e['year']} | Evidence: {mask_evidence_id(e['evidence_id'])}")
                    st.write(f"  - Generated Column 3 preview (first 200 chars): {e['generated_statement'][:200]}...")
            else:
                st.write("_No qualifications entered yet._")
            st.write("**Part 2 (Role):**")
            if u["part2"]:
                st.write(f"- Role: {u['part2']['role_title']} | Years: {u['part2']['years_exp']} | Evidence: {mask_evidence_id(u['part2']['evidence_id'])}")
                st.write(f"  - Column 3 preview: {u['part2']['generated_statement'][:200]}...")
            else:
                st.write("_No role entered yet._")
            st.write("**Part 3 (PD):**")
            if u["part3"]:
                st.write(f"- PD: {u['part3']['pd_title']} | Year: {u['part3']['year']} | Evidence: {mask_evidence_id(u['part3']['evidence_id'])}")
                st.write(f"  - Column 3 preview: {u['part3']['generated_statement'][:200]}...")
            else:
                st.write("_No PD entered yet._")

        show_unit_summary(st.session_state.units_data[key])
        col_confirm = st.columns(3)
        if col_confirm[0].button(f"Mark {unit_code} as Complete", key=f"confirm_{unit_code}"):
            st.session_state.units_data[key]["confirmed"] = True
            st.success(f"{unit_code} marked complete. You can still edit from the left expander if needed.")
        if col_confirm[1].button(f"Reset entries for {unit_code}", key=f"reset_{unit_code}"):
            st.session_state.units_data[key] = {
                "unit_code": unit_code,
                "unit_name": unit_name,
                "part1": [],
                "part2": None,
                "part3": None,
                "confirmed": False
            }
            st.experimental_rerun()

# ---------- QA & Export ----------
st.header("Step 4 — QA, Edit, and Export")
# Build a flat list of all rows that will be inserted (Part1 rows first, then part2 and part3)
all_rows = []
for key, data in st.session_state.units_data.items():
    ucode = data["unit_code"]
    uname = data["unit_name"]
    # Part 1 rows
    for p in data["part1"]:
        row = {
            "unit_code": ucode,
            "col1": p["qual_name"],
            "col2": p["year"],
            "col3": p["generated_statement"],
            "col4": p["evidence_id"]
        }
        all_rows.append(row)
    # Part 2 single row (if provided)
    if data.get("part2") and (data["part2"].get("role_title") or data["part2"].get("evidence_id")):
        r = data["part2"]
        row = {"unit_code": ucode, "col1": r["role_title"], "col2": r["years_exp"], "col3": r["generated_statement"], "col4": r["evidence_id"]}
        all_rows.append(row)
    # Part 3 single row
    if data.get("part3") and (data["part3"].get("pd_title") or data["part3"].get("evidence_id")):
        r = data["part3"]
        row = {"unit_code": ucode, "col1": r["pd_title"], "col2": r["year"], "col3": r["generated_statement"], "col4": r["evidence_id"]}
        all_rows.append(row)

st.write(f"Total rows to insert: {len(all_rows)}")
if not all_rows:
    st.info("No rows collected yet. Complete at least one unit to export.")
    st.stop()

st.subheader("QA Summary — masked Evidence IDs shown")
for i, row in enumerate(all_rows, start=1):
    st.write(f"Row {i}: {row['unit_code']} | {row['col1']} | {row['col2']} | Evidence: {mask_evidence_id(row['col4'])}")

# Final checks
invalid_years = [r for r in all_rows if r['col2'] and not validate_year(r['col2'])]
if invalid_years:
    st.error("Some rows have invalid years. Please edit and fix them before export.")
    st.stop()

# Show Pending IDs
pending = [r for r in all_rows if (not r['col4']) or r['col4'].strip().lower() == "pending"]
if pending:
    st.warning(f"{len(pending)} rows have 'Pending' or missing Evidence ID. They will be flagged in the exported file.")

# Allow final edit toggle (simple: let them re-open unit expanders above)
st.write("If you need to edit entries, open the unit expanders above and press 'Reset' or modify fields. Once ready, export the VCMT.")

# Export action - write into a copy of the uploaded doc
if st.button("Generate and Download VCMT (.docx)"):
    out_doc = load_docx(uploaded.read())  # reload fresh
    # find part1 table index
    t_index = find_part1_table_index(out_doc)
    if t_index == -1:
        # append new table at end with 4 columns
        t = out_doc.add_table(rows=1, cols=4)
        hdr_cells = t.rows[0].cells
        hdr_cells[0].text = "Column 1"
        hdr_cells[1].text = "Column 2"
        hdr_cells[2].text = "Column 3"
        hdr_cells[3].text = "Column 4"
        t_index = len(out_doc.tables) - 1
    table = out_doc.tables[t_index]

    # Append each collected row
    for r in all_rows:
        new_row = table.add_row()
        # Column 1: name/title
        new_row.cells[0].text = r["col1"]
        # Column 2: year / years
        new_row.cells[1].text = r["col2"] or ""
        # Column 3: generated statement (write the full generated text)
        new_row.cells[2].text = r["col3"] or ""
        # Column 4: full Evidence ID (not masked)
        new_row.cells[3].text = r["col4"] or ""

    # QA: insert a simple flag row if pending IDs exist
    if pending:
        note_row = table.add_row()
        note_row.cells[0].text = "NOTE"
        note_row.cells[1].text = ""
        note_row.cells[2].text = f"{len(pending)} row(s) have Pending or missing Evidence IDs. Please update."
        note_row.cells[3].text = ""

    # Build filename
    today = datetime.date.today().isoformat().replace("-", "")
    codes_for_name = "_".join(validated_codes)
    filename = f"VCMT_{codes_for_name}_{today}.docx"

    bio = io.BytesIO()
    out_doc.save(bio)
    bio.seek(0)
    st.success("VCMT generated.")
    st.download_button("Download filled VCMT (.docx)", data=bio.read(), file_name=filename, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
